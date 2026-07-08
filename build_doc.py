"""
Convert SIMAX_QUANTA_MODULE_GUIDE.md to a formatted Word document.
Run with: backend/venv/Scripts/python.exe build_doc.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE  = "SIMAX_QUANTA_MODULE_GUIDE.md"
OUT_FILE = "Simax_Quanta_Module_Guide.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1E, 0x3A, 0x5F)   # headings
BLUE_MID   = RGBColor(0x2E, 0x6D, 0xB4)   # h2
BLUE_LIGHT = RGBColor(0x37, 0x8B, 0xC0)   # h3
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)   # body
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TEAL       = RGBColor(0x1A, 0x73, 0x6B)   # code bg header
TABLE_HDR  = RGBColor(0x1E, 0x3A, 0x5F)


def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DB4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_inline(run_parent, text: str):
    """Parse **bold** and `code` inline markers and add runs to paragraph."""
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = run_parent.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = GRAY_TEXT
        elif part.startswith("`") and part.endswith("`"):
            run = run_parent.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            if part:
                run = run_parent.add_run(part)
                run.font.color.rgb = GRAY_TEXT


def style_doc(doc):
    """Set document-wide default font and margins."""
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = GRAY_TEXT


def add_title_page(doc, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(13)
    r2.font.color.rgb = BLUE_MID
    r2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Version 4.0  |  Confidential — Internal Use Only")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r3.font.italic = True

    doc.add_page_break()


def add_h1(doc, text):
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE_MID
    run.font.name = "Calibri"


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_LIGHT
    run.font.name = "Calibri"


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_DARK
    run.font.italic = True


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    apply_inline(p, text)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    apply_inline(p, text)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(doc, rows):
    if not rows:
        return
    # Detect header separator row (---|---|---)
    data_rows = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if len(data_rows) < 1:
        return

    def parse_row(row):
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        return cells

    parsed = [parse_row(r) for r in data_rows]
    col_count = max(len(r) for r in parsed)

    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, row_data in enumerate(parsed):
        row_obj = table.rows[ri]
        for ci in range(col_count):
            cell = row_obj.cells[ci]
            text = row_data[ci] if ci < len(row_data) else ""

            # Clean markdown bold from header
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)

            cell.text = text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

            if ri == 0:
                run.bold = True
                run.font.color.rgb = WHITE
                set_cell_bg(cell, "1E3A5F")
            else:
                run.font.color.rgb = GRAY_TEXT
                bg = "FFFFFF" if ri % 2 == 1 else "EEF4FB"
                set_cell_bg(cell, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Main parser ────────────────────────────────────────────────────────────────

def convert(md_path: str, out_path: str):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    style_doc(doc)
    add_title_page(doc, "Simax Quanta", "Government Knowledge Intelligence Platform — Module Guide")

    in_code  = False
    in_table = False
    code_buf = []
    tbl_buf  = []

    for raw in lines:
        line = raw.rstrip("\n")

        # ── Code block toggle ────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code  = False
            else:
                if in_table:
                    add_table(doc, tbl_buf)
                    tbl_buf  = []
                    in_table = False
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Table rows ───────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            in_table = True
            tbl_buf.append(line)
            continue
        else:
            if in_table:
                add_table(doc, tbl_buf)
                tbl_buf  = []
                in_table = False

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        if line.startswith("#### "):
            add_h4(doc, line[5:].strip())
        elif line.startswith("### "):
            add_h3(doc, line[4:].strip())
        elif line.startswith("## "):
            add_h2(doc, line[3:].strip())
        elif line.startswith("# "):
            add_h1(doc, line[2:].strip())

        # ── Bullet lists ─────────────────────────────────────────────────────
        elif re.match(r'^(\s*)[-*] ', line):
            indent = len(line) - len(line.lstrip())
            level  = indent // 2
            text   = re.sub(r'^\s*[-*] ', '', line)
            add_bullet(doc, text.strip(), level)

        # ── Numbered lists ───────────────────────────────────────────────────
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            apply_inline(p, text)

        # ── Blank line ───────────────────────────────────────────────────────
        elif line.strip() == "":
            continue

        # ── Regular paragraph ────────────────────────────────────────────────
        else:
            add_body(doc, line.strip())

    # Flush any remaining buffers
    if in_code and code_buf:
        add_code_block(doc, code_buf)
    if in_table and tbl_buf:
        add_table(doc, tbl_buf)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
