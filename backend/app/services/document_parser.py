"""
Module 1 — Document Ingestion Engine
Module 2 — Structure-Aware Chunking

Fixes over previous version:
- Sentence-aware splitting: never cuts inside a sentence or list point
- Abbreviation protection: Rs., No., Govt., Sec., e.g., w.e.f. not treated as sentence ends
- Improved heading detection: ALL-CAPS, numbered sections (1., A., i.), underlined text
- List preservation: bullet/numbered points joined with newlines, never split mid-list
- Larger chunk window: 2000 chars to better use BGE Large's 512-token capacity
- Sentence-level overlap: last 2 sentences carried to next chunk for context continuity
- XLSX: column header row pinned to every row-group chunk

Supported: PDF (text + scanned/OCR fallback), XLSX, DOCX, TXT
Returns: (extracted_text: str, chunks: list[dict])
Each chunk: { chunk_text, heading_path, chunk_type }
"""
import io
import re
import csv
import fitz
import openpyxl
from docx import Document as DocxDocument
from dotenv import load_dotenv
import os

load_dotenv()

# Auto-detect Tesseract from env var or common install paths
_TESSERACT_CMD = os.getenv("TESSERACT_CMD", "")
if not _TESSERACT_CMD:
    for _candidate in [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
    ]:
        if os.path.exists(_candidate):
            _TESSERACT_CMD = _candidate
            break

try:
    import pytesseract
    from PIL import Image
    if _TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = _TESSERACT_CMD
    _OCR_AVAILABLE = True
except Exception:
    _OCR_AVAILABLE = False

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".txt", ".docx", ".csv"}

MAX_CHUNK_CHARS    = 2000   # was 1200 — BGE Large handles 512 tokens ≈ 3000+ chars
MIN_CHUNK_CHARS    = 40
MAX_ROWS_PER_CHUNK = 40
_OVERLAP_SENTS     = 2      # carry last N sentences into next chunk for continuity


# ── Text helpers ──────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Strip control chars, collapse blank lines and spaces."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


_WATERMARK_RE = re.compile(
    r'(look deep into nature|albert einstein|smart and effective notes|upsc cse|'
    r'visit.*\.com|www\.|copyright|all rights reserved|page \d+|^\d+$)',
    re.IGNORECASE
)

def _is_watermark(text: str) -> bool:
    return bool(_WATERMARK_RE.search(text)) or len(text) > 300

def _heading_path(stack: list) -> str:
    return " > ".join(t for _, t in stack if not _is_watermark(t))

def _median_size(blocks: list) -> float:
    sizes = [
        span["size"]
        for b in blocks if b.get("type") == 0
        for line in b.get("lines", [])
        for span in line.get("spans", [])
        if span.get("text", "").strip()
    ]
    if not sizes:
        return 12.0
    sizes.sort()
    return sizes[len(sizes) // 2]


# ── Sentence-aware splitter ───────────────────────────────────────────────────

# Abbreviations common in Indian government documents — period after these is NOT a sentence end
_ABBREV_RE = re.compile(
    r'\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|St|No|Nos|Rs|Sh|Smt|vs|etc|e\.g|i\.e|'
    r'approx|Dept|Govt|Sec|Art|Vol|pg|pp|Ed|Fig|viz|w\.e\.f|w\.r\.t|'
    r'Cl|Sub|para|Para|Ref|Sl|S\.No|S\.No|Jan|Feb|Mar|Apr|Jun|Jul|Aug|'
    r'Sep|Oct|Nov|Dec)\.',
    re.IGNORECASE
)

# List item markers at line start — each is an atomic unit, never split inside
_LIST_ITEM_RE = re.compile(
    r'^[\s]*([•\-\*◦▪➢➤►→]|\d+[\.\)]\s|\(?[a-zA-Z]\)?[\.\)]\s|[ivxlIVX]+[\.\)]\s)'
)

def _is_list_item(text: str) -> bool:
    return bool(_LIST_ITEM_RE.match(text.strip()))


def _split_sentences(text: str) -> list:
    """
    Split text into sentence/item units without breaking on abbreviations.
    Multi-line text (lists) is split line-by-line so each item stays intact.
    """
    lines = text.split('\n')
    if len(lines) > 1:
        sentences = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Keep list items atomic — never split inside them
            if _is_list_item(line):
                sentences.append(line)
            else:
                sentences.extend(_split_sentences_flat(line))
        return sentences
    return _split_sentences_flat(text)


def _split_sentences_flat(text: str) -> list:
    """Split a flat (single-line) string into sentences."""
    # Temporarily mark abbreviation periods so they survive the split
    protected = _ABBREV_RE.sub(lambda m: m.group().rstrip('.') + '.\x01', text)
    # Split at real sentence boundaries: end-punctuation + space + capital/quote
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z\"\'\(])', protected)
    result = [p.replace('\x01', ' ').strip() for p in parts if p.strip()]
    return result if result else [text.strip()]


def _pack_sentences(sentences: list, heading_path: str, chunk_type: str) -> list:
    """
    Pack sentences into chunks ≤ MAX_CHUNK_CHARS.
    Carries last _OVERLAP_SENTS sentences into the next chunk for context.
    Never breaks a sentence or list item in half.
    """
    chunks  = []
    current: list = []
    current_len   = 0

    for sent in sentences:
        sent_len = len(sent) + 1  # +1 for newline separator

        # Single sentence too big on its own — emit alone
        if len(sent) > MAX_CHUNK_CHARS:
            if current:
                chunks.append(_make_chunk('\n'.join(current), heading_path, chunk_type))
                current = current[-_OVERLAP_SENTS:]
                current_len = sum(len(s) + 1 for s in current)
            chunks.append(_make_chunk(sent, heading_path, chunk_type))
            continue

        # Current chunk would overflow — flush and start new with overlap
        if current_len + sent_len > MAX_CHUNK_CHARS and current:
            chunks.append(_make_chunk('\n'.join(current), heading_path, chunk_type))
            current = current[-_OVERLAP_SENTS:]
            current_len = sum(len(s) + 1 for s in current)

        current.append(sent)
        current_len += sent_len

    if current:
        text = '\n'.join(current).strip()
        if len(text) >= MIN_CHUNK_CHARS:
            chunks.append(_make_chunk(text, heading_path, chunk_type))

    return chunks


def _make_chunk(text: str, heading_path: str, chunk_type: str) -> dict:
    return {"chunk_text": text.strip(), "heading_path": heading_path, "chunk_type": chunk_type}


def _emit_paragraph(parts: list, heading_path: str, chunks: list) -> None:
    """
    Merge accumulated paragraph blocks and emit as sentence-packed chunks.
    List items are joined with newlines to preserve structure.
    Plain paragraphs are joined with spaces.
    """
    if not parts:
        return

    has_list = any(_is_list_item(p) for p in parts)
    text = '\n'.join(parts).strip() if has_list else ' '.join(parts).strip()

    if not text or len(text) < MIN_CHUNK_CHARS:
        return

    sentences = _split_sentences(text)
    chunks.extend(_pack_sentences(sentences, heading_path, "paragraph"))


# ── PDF heading detection ─────────────────────────────────────────────────────

# Numbered section patterns common in Indian government documents
_NUMBERED_SECTION_RE = re.compile(
    r'^(\d+\.\s|\d+\.\d+[\s\.]+|[A-Z]\.\s|[IVX]+\.\s|\([a-z]\)\s|\([ivx]+\)\s)'
)

def _is_pdf_heading(block_text: str, max_size: float, body_size: float,
                    all_bold: bool, all_underline: bool) -> bool:
    """
    Multi-signal heading detection for Indian government PDFs.
    Covers font size, bold, ALL CAPS, numbered sections, underline.
    """
    text   = block_text.strip()
    length = len(text)

    # 1. Font significantly larger than body text
    if max_size > body_size + 1.5:
        return True

    # 2. Bold AND short — classic heading pattern
    if all_bold and length < 120 and '\n' not in text:
        return True

    # 3. ALL CAPS with multiple words — section headings like "OBJECTIVES", "ELIGIBILITY CRITERIA"
    if text.isupper() and length > 4 and length < 120 and ' ' in text:
        return True

    # 4. Underlined text — heading style in older government circulars/GOs
    if all_underline and length < 150:
        return True

    # 5. Numbered section: "1. ", "2.1 ", "A. ", "i. " followed by capital letter
    if _NUMBERED_SECTION_RE.match(text) and length < 150:
        return True

    return False


# ── PDF parser ────────────────────────────────────────────────────────────────

def _parse_pdf(path: str):
    doc           = fitz.open(path)
    chunks        = []
    full_text     = []
    heading_stack = []
    para_parts    = []

    def flush():
        _emit_paragraph(para_parts, _heading_path(heading_stack), chunks)
        para_parts.clear()

    for page_num, page in enumerate(doc):
        page_dict = page.get_text("dict")
        blocks    = page_dict.get("blocks", [])
        body_size = _median_size(blocks)
        page_text = []

        for block in blocks:
            if block.get("type") != 0:
                continue

            block_parts   = []
            line_texts    = []   # one entry per line — preserves list structure
            max_size      = 0.0
            all_bold      = True
            all_underline = True

            for line in block.get("lines", []):
                line_parts = []
                for span in line.get("spans", []):
                    t = span.get("text", "").strip()
                    if not t:
                        continue
                    block_parts.append(t)
                    line_parts.append(t)
                    size  = span.get("size", 12)
                    flags = span.get("flags", 0)
                    max_size = max(max_size, size)
                    if not (flags & 16):  # bit 4 = bold
                        all_bold = False
                    if not (flags & 4):   # bit 2 = underline
                        all_underline = False
                if line_parts:
                    line_texts.append(" ".join(line_parts))

            if not block_parts:
                continue

            block_text = " ".join(block_parts).strip()
            if not block_text or len(block_text) < MIN_CHUNK_CHARS:
                continue

            page_text.append(block_text)

            if _is_pdf_heading(block_text, max_size, body_size, all_bold, all_underline):
                if _is_watermark(block_text):
                    continue
                flush()
                # Use a virtual size for non-font-size headings so hierarchy still works
                h_size = max_size if max_size > body_size + 1.5 else body_size + 0.5
                heading_stack = [(s, t) for s, t in heading_stack if s > h_size]
                heading_stack.append((h_size, block_text))
                chunks.append(_make_chunk(block_text, _heading_path(heading_stack), "heading"))
            else:
                # Detect list blocks: multiple lines where some are list items
                if len(line_texts) > 1 and any(_is_list_item(l) for l in line_texts):
                    para_parts.extend(l for l in line_texts if l.strip())
                else:
                    para_parts.append(block_text)

                if sum(len(p) for p in para_parts) >= MAX_CHUNK_CHARS:
                    flush()

        # OCR fallback for scanned/image-only pages
        if not page_text and _OCR_AVAILABLE:
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr = pytesseract.image_to_string(img, lang="eng", config="--oem 3 --psm 6")
                if ocr.strip():
                    page_text.append(ocr)
                    ocr_cleaned = _clean(ocr)
                    if len(ocr_cleaned) >= MIN_CHUNK_CHARS:
                        chunks.extend(_pack_sentences(
                            _split_sentences(ocr_cleaned),
                            _heading_path(heading_stack),
                            "paragraph"
                        ))
            except Exception as e:
                print(f"[OCR] page {page_num + 1}: {e}")

        full_text.extend(page_text)

    flush()
    doc.close()
    return _clean("\n".join(full_text)), chunks


# ── XLSX parser ───────────────────────────────────────────────────────────────

def _parse_xlsx(path: str):
    wb = openpyxl.load_workbook(path, data_only=True)
    chunks, all_text = [], []

    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))

        if not rows:
            continue

        sheet_header = f"[Sheet: {sheet.title}]"
        all_text.append(sheet_header)
        all_text.extend(rows)

        # Pin column header row to every chunk so each chunk is self-contained
        header_row = rows[0]
        data_rows  = rows[1:] if len(rows) > 1 else []

        if not data_rows:
            chunks.append(_make_chunk(f"{sheet_header}\n{header_row}", sheet.title, "table"))
            continue

        for i in range(0, len(data_rows), MAX_ROWS_PER_CHUNK):
            group      = data_rows[i: i + MAX_ROWS_PER_CHUNK]
            chunk_text = f"{sheet_header}\n{header_row}\n" + "\n".join(group)
            chunks.append(_make_chunk(chunk_text, sheet.title, "table"))

    wb.close()
    return _clean("\n\n".join(all_text)), chunks


# ── TXT parser ────────────────────────────────────────────────────────────────

def _parse_txt(path: str):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    cleaned = _clean(text)

    raw_paras = [p.strip() for p in re.split(r"\n\s*\n", cleaned) if p.strip()]

    chunks    = []
    merge_buf = []
    merge_len = 0

    for para in raw_paras:
        if len(para) < MIN_CHUNK_CHARS:
            merge_buf.append(para)
            merge_len += len(para)
            if merge_len >= MIN_CHUNK_CHARS:
                merged = '\n'.join(merge_buf).strip()
                chunks.extend(_pack_sentences(_split_sentences(merged), "", "paragraph"))
                merge_buf.clear()
                merge_len = 0
        else:
            if merge_buf:
                merged = '\n'.join(merge_buf).strip()
                if len(merged) >= MIN_CHUNK_CHARS:
                    chunks.extend(_pack_sentences(_split_sentences(merged), "", "paragraph"))
                merge_buf.clear()
                merge_len = 0
            chunks.extend(_pack_sentences(_split_sentences(para), "", "paragraph"))

    if merge_buf:
        merged = '\n'.join(merge_buf).strip()
        if len(merged) >= MIN_CHUNK_CHARS:
            chunks.extend(_pack_sentences(_split_sentences(merged), "", "paragraph"))

    return cleaned, chunks


# ── DOCX parser ───────────────────────────────────────────────────────────────

_DOCX_HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "heading 4", "heading 5"}

def _parse_docx(path: str):
    doc           = DocxDocument(path)
    chunks        = []
    full_text     = []
    heading_stack = []
    para_parts    = []

    def _docx_path(stack):
        return " > ".join(t for _, t in stack)

    def flush():
        _emit_paragraph(para_parts, _docx_path(heading_stack), chunks)
        para_parts.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style and para.style.name else ""
        is_heading = style_name in _DOCX_HEADING_STYLES
        level      = int(style_name[-1]) if is_heading else 0

        # Fallback: bold-only short paragraphs treated as headings (unstyle docs)
        if not is_heading and para.runs and len(text) < 120:
            if all(run.bold for run in para.runs if run.text.strip()):
                is_heading = True
                level = 3

        # Fallback: ALL CAPS short line
        if not is_heading and text.isupper() and 4 < len(text) < 120 and ' ' in text:
            is_heading = True
            level = 2

        full_text.append(text)

        if is_heading:
            flush()
            heading_stack = [(l, t) for l, t in heading_stack if l < level]
            heading_stack.append((level, text))
            if len(text) >= MIN_CHUNK_CHARS:
                chunks.append(_make_chunk(text, _docx_path(heading_stack), "heading"))
        else:
            para_parts.append(text)
            if sum(len(p) for p in para_parts) >= MAX_CHUNK_CHARS:
                flush()

    # Tables — pin header row to every chunk
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if not rows:
            continue
        full_text.extend(rows)
        header_row = rows[0]
        data_rows  = rows[1:] if len(rows) > 1 else []
        if not data_rows:
            chunks.append(_make_chunk(header_row, _docx_path(heading_stack), "table"))
            continue
        for i in range(0, len(data_rows), MAX_ROWS_PER_CHUNK):
            group = data_rows[i: i + MAX_ROWS_PER_CHUNK]
            table_text = header_row + "\n" + "\n".join(group)
            chunks.extend(_pack_sentences(_split_sentences(table_text), _docx_path(heading_stack), "table"))

    flush()
    return _clean("\n".join(full_text)), chunks


# ── CSV parser ────────────────────────────────────────────────────────────────

def _parse_csv(path: str):
    chunks, all_text = [], []

    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        rows = [row for row in reader if any(cell.strip() for cell in row)]

    if not rows:
        return "", []

    header_row = " | ".join(str(c).strip() for c in rows[0])
    all_text.append("[CSV File]")
    all_text.append(header_row)

    data_rows = [
        " | ".join(str(c).strip() for c in row)
        for row in rows[1:]
        if any(str(c).strip() for c in row)
    ]
    all_text.extend(data_rows)

    if not data_rows:
        chunks.append(_make_chunk(f"[CSV]\n{header_row}", "", "table"))
    else:
        for i in range(0, len(data_rows), MAX_ROWS_PER_CHUNK):
            group = data_rows[i: i + MAX_ROWS_PER_CHUNK]
            chunk_text = f"[CSV]\n{header_row}\n" + "\n".join(group)
            chunks.append(_make_chunk(chunk_text, "", "table"))

    return _clean("\n".join(all_text)), chunks


# ── Public API ────────────────────────────────────────────────────────────────

def parse_document(path: str):
    ext = ("." + path.rsplit(".", 1)[-1].lower()) if "." in path else ""
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".xlsx":
        return _parse_xlsx(path)
    if ext == ".csv":
        return _parse_csv(path)
    if ext == ".txt":
        return _parse_txt(path)
    if ext == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported file type: '{ext}'")
